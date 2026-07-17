import streamlit as st
import pandas as pd
from streamlit_calendar import calendar
from datetime import datetime, time, date
import re
import os
import io
import requests
import base64

# Importación segura de la librería python-docx
try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    LIBRERIA_DOCX_LISTA = True
except ImportError:
    LIBRERIA_DOCX_LISTA = False

# ==========================================
# 1. CONFIGURACIÓN DE LA PÁGINA Y ESTILOS (IDENTIDAD VISUAL OFICIAL)
# ==========================================
st.set_page_config(
    layout="wide", 
    page_title="Calendario territorial UPEU", 
    page_icon="🗓️"
)

# Inyección de estilos CSS basados estrictamente en el Manual de Marca de Río Negro
st.markdown("""
    <style>
    /* Importación de la tipografía oficial Figtree */
    @import url('https://fonts.googleapis.com/css2?family=Figtree:ital,wght@0,300..900;1,300..900&display=swap');
    
    /* Configuración de fuentes globales */
    html, body, [class*="css"], .stMarkdown, p, div {
        font-family: 'Figtree', sans-serif !important;
    }
    
    /* Fondo general utilizando el Gris RN oficial (#E8E8E8) */
    .stApp, .main, [data-testid="stAppViewContainer"], [data-testid="stHeader"], [data-testid="stMainBlockContainer"] { 
        background-color: #E8E8E8 !important; 
        background: #E8E8E8 !important;
    }
    
    /* Diseño de tarjetas y formularios (Fondo blanco rígido y bordes limpios) */
    div[data-testid="stForm"] { 
        background-color: #FFFFFF !important; 
        border-radius: 8px !important; 
        padding: 30px !important; 
        box-shadow: 0px 4px 6px rgba(0, 0, 0, 0.05) !important;
        border: 1px solid #D1D5DB !important;
    }
    
    /* Títulos Principales en Negro Puro */
    h1 { 
        color: #000000 !important; 
        font-weight: 800 !important;
        font-size: 2.2rem !important;
        margin-top: 10px !important;
        margin-bottom: 0px !important;
    }
    
    /* Subtítulos en el Azul RN oficial (#007BE0) */
    h2, h3 { 
        color: #007BE0 !important; 
        font-weight: 700 !important;
    }
    
    /* Estilo del botón primario con el Verde RN oficial (#6AC64F) */
    button[kind="primary"] {
        background-color: #6AC64F !important;
        color: #FFFFFF !important;
        font-weight: 700 !important;
        border-radius: 6px !important;
        border: none !important;
        transition: background-color 0.2s ease;
    }
    
    button[kind="primary"]:hover {
        background-color: #59b040 !important; /* Verde RN ligeramente más oscuro para el hover */
    }
    
    /* Hashtag de gestión oficial #gobiernodelosrionegrinos */
    .hashtag-gestion {
        color: #6AC64F !important;
        font-weight: 800;
        font-size: 1.1rem;
    }

    /* Forzar nitidez absoluta en el renderizado de imágenes vectoriales */
    img {
        image-rendering: -webkit-optimize-contrast !important;
        image-rendering: crisp-edges !important;
    }

    /* ==========================================
       ANULAR EL ROJO NATIVO DE STREAMLIT (TABS Y SELECCIONES)
       ========================================== */
    div[data-baseweb="tab-list"] button[aria-selected="true"] {
        color: #007BE0 !important;
        border-bottom-color: #007BE0 !important;
    }
    div[data-baseweb="tab-list"] button[aria-selected="false"] {
        color: #333333 !important;
    }
    .stTextInput input:focus, 
    .stSelectbox div[role="button"]:focus, 
    .stTextArea textarea:focus,
    div[data-baseweb="select"] > div:focus-within {
        border-color: #007BE0 !important;
        box-shadow: 0 0 0 1px #007BE0 !important;
    }
    div[data-testid="stSpinner"] > div {
        border-top-color: #6AC64F !important;
    }

    /* ==========================================
       PERSONALIZACIÓN CSS PARA EL CALENDARIO (FULLCALENDAR)
       ========================================== */
    .fc .fc-button,
    .fc .fc-button-primary,
    .fc-button,
    .fc-button-primary,
    button.fc-button,
    button.fc-today-button,
    button.fc-prev-button,
    button.fc-next-button {
        background-color: #007BE0 !important;
        background: #007BE0 !important;
        border-color: #007BE0 !important;
        color: #FFFFFF !important;
        opacity: 1 !important;
        box-shadow: none !important;
        font-family: 'Figtree', sans-serif !important;
        font-weight: 600 !important;
        text-transform: capitalize !important;
        transition: background-color 0.2s ease, border-color 0.2s ease !important;
    }
    .fc .fc-button:hover,
    .fc .fc-button-primary:hover,
    .fc-button:hover,
    .fc-button-primary:hover,
    button.fc-button:hover,
    button.fc-today-button:hover,
    button.fc-prev-button:hover,
    button.fc-next-button:hover {
        background-color: #6AC64F !important;
        background: #6AC64F !important;
        border-color: #6AC64F !important;
        color: #FFFFFF !important;
    }
    .fc .fc-button-primary:not(:disabled).fc-button-active, 
    .fc .fc-button-primary:not(:disabled):active,
    .fc-button-active,
    button.fc-button-active,
    .fc .fc-button-primary:active {
        background-color: #00569E !important;
        background: #00569E !important;
        border-color: #00569E !important;
        color: #FFFFFF !important;
    }
    .fc .fc-col-header-cell-cushion {
        color: #000000 !important;
        font-weight: 700 !important;
        text-decoration: none !important;
    }
    .fc-event, .fc-event-dot {
        border-color: transparent !important;
    }
    .fc .fc-daygrid-day.fc-day-today {
        background-color: rgba(0, 123, 224, 0.08) !important;
    }
    .fc .fc-day-today .fc-daygrid-day-number {
        color: #007BE0 !important;
        font-weight: 800 !important;
    }
    .fc .fc-day-today .fc-daygrid-day-top {
        border-top: 3px solid #007BE0 !important;
    }
    .fc .fc-timegrid-now-indicator-line {
        border-color: #6AC64F !important;
        border-width: 2px !important;
    }
    .fc .fc-timegrid-now-indicator-arrow {
        border-top-color: #6AC64F !important;
        border-bottom-color: #6AC64F !important;
    }
    </style>
""", unsafe_allow_html=True)

LOGO_FILE = "isologo_RN.svg"

# ==========================================
# 2. CONTROL DE ACCESO (MODO LECTOR / EDITOR)
# ==========================================
st.sidebar.header("🔑 Control de Acceso")
password = st.sidebar.text_input("Contraseña de Editor", type="password")

CONTRASEÑA_CORRECTA = "UPEU2026" 
es_editor = (password == CONTRASEÑA_CORRECTA)

if es_editor:
    st.sidebar.success("🔑 Modo Editor Activado")
else:
    st.sidebar.info("👁️ Modo Visualización (Solo Lectura)")

# ==========================================
# 3. FUNCIONES DE PERSISTENCIA DIRECTA CON GITHUB API
# ==========================================

def github_request(method, payload=None):
    token = st.secrets.get("GITHUB_TOKEN")
    repo = st.secrets.get("GITHUB_REPO")
    path = st.secrets.get("GITHUB_FILE_PATH", "agenda_territorial_consolidada.xlsx")
    
    if not token or not repo:
        return None, "Faltan configurar las credenciales de GitHub en los Secrets de Streamlit."
        
    url = f"https://api.github.com/repos/{repo}/contents/{path}"
    headers = {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github.v3+json"
    }
    
    if method == "GET":
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            return response.json(), None
        return None, f"Error al descargar archivo de GitHub (Status: {response.status_code})"
        
    elif method == "PUT" and payload:
        response = requests.put(url, headers=headers, json=payload)
        if response.status_code in [200, 201]:
            return response.json(), None
        return None, f"Error al guardar cambios en GitHub: {response.text}"
        
    return None, "Método inválido"

def limpiar_fecha_para_calendario(val):
    val_str = str(val).strip()
    if not val_str or val_str == "nan" or val_str == "":
        return None
    match_iso = re.match(r"^(\d{4}-\d{2}-\d{2})", val_str)
    if match_iso:
        return match_iso.group(1)
    val_str = re.sub(r"\s*/\s*", "/", val_str)
    match_rango = re.search(r"(\d+)\s*(?:y|a|-)\s*\d+/(\d+)/(\d{3,4})", val_str)
    if match_rango:
        return f"{match_rango.group(3)}-{match_rango.group(2).zfill(2)}-{match_rango.group(1).zfill(2)}"
    match_normal = re.match(r"^(\d{1,2})/(\d{1,2})/(\d{3,4})", val_str)
    if match_normal:
        return f"{match_normal.group(3)}-{match_normal.group(2).zfill(2)}-{match_normal.group(1).zfill(2)}"
    return None

def obtener_mes_nombre(val_fecha):
    meses_es = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    f_limpia = limpiar_fecha_para_calendario(val_fecha)
    if f_limpia:
        try:
            return meses_es[int(f_limpia.split("-")[1])]
        except:
            pass
    for m in meses_es:
        if m and m.lower() in str(val_fecha).lower():
            return m
    return "Fecha Flexible / Por definir"

def load_data():
    file_info, error = github_request("GET")
    columnas_requeridas = [
        'Fecha', 'Hora', 'Semana', 'Actividad', 'Ciudad', 'Lugar',
        'Explicación breve de la actividad', 'Cantidad de personas estimadas',
        'Organismo/Actor', 'Estado', 'Público Destinatario', 'Prioridad', 'Invitación a participar'
    ]
    
    if error or not file_info:
        path_local = st.secrets.get("GITHUB_FILE_PATH", "agenda_territorial_consolidada.xlsx")
        if os.path.exists(path_local):
            df = pd.read_excel(path_local)
        else:
            return pd.DataFrame(columns=columnas_requeridas)
    else:
        conte_bytes = base64.b64decode(file_info["content"])
        df = pd.read_excel(io.BytesIO(conte_bytes))
        
    df.columns = df.columns.str.strip()
    
    df['_temp_f'] = df['Fecha'].astype(str).str.strip().replace('', None)
    df['_temp_a'] = df['Actividad'].astype(str).str.strip().replace('', None)
    df = df[(~df['_temp_f'].isna() & (df['_temp_f'] != 'nan')) | (~df['_temp_a'].isna() & (df['_temp_a'] != 'nan'))]
    df = df.drop(columns=['_temp_f', '_temp_a'])

    for col in columnas_requeridas:
        if col not in df.columns:
            df[col] = ""
            
    df['Actividad'] = df['Actividad'].fillna("Actividad sin título").astype(str)
    df['Hora'] = df['Hora'].fillna("Sin especificar").astype(str).str.strip()
    df['Ciudad'] = df['Ciudad'].fillna("Sin Ciudad").astype(str).str.strip()
    df['Lugar'] = df['Lugar'].fillna("Sin especificar").astype(str).str.strip()
    df['Explicación breve de la actividad'] = df['Explicación breve de la actividad'].fillna("").astype(str)
    df['Cantidad de personas estimadas'] = pd.to_numeric(df['Cantidad de personas estimadas'], errors='coerce').fillna(0).astype(int)
    df['Organismo/Actor'] = df['Organismo/Actor'].fillna("No especificado").astype(str)
    df['Estado'] = df['Estado'].fillna("Pendiente").astype(str)
    df['Prioridad'] = df['Prioridad'].fillna("INTERMEDIA").astype(str)
    df['Público Destinatario'] = df['Público Destinatario'].fillna("General").astype(str)
    df['Invitación a participar'] = df['Invitación a participar'].fillna("").astype(str)
    
    return df.reset_index(drop=True)

def push_data_to_github(df, commit_message="Actualización desde el panel de control territorial"):
    file_info, error = github_request("GET")
    if error or not file_info:
        st.error(f"No se pudo sincronizar el estado previo en GitHub: {error}")
        return False
        
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    excel_bytes = output.getvalue()
    content_b64 = base64.b64encode(excel_bytes).decode("utf-8")
    
    payload = {
        "message": commit_message,
        "content": content_b64,
        "sha": file_info["sha"]
    }
    
    _, put_error = github_request("PUT", payload=payload)
    if put_error:
        st.error(put_error)
        return False
    return True

if 'agenda' not in st.session_state:
    st.session_state.agenda = load_data()

# ==========================================
# EXPORTACIÓN A WORD Y WHATSAPP
# ==========================================

def set_cell_background(cell, color_hex):
    try:
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
    except:
        pass

def crear_reporte_word_areas(df, titulo_personalizado="REPORTE PLANIFICACION TERRITORIAL", aclaracion_rango=""):
    doc = docx.Document()
    try:
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    except:
        pass
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    p_header = doc.add_paragraph()
    run_gob = p_header.add_run("GOBIERNO DE LA PROVINCIA DE RÍO NEGRO\n")
    run_gob.font.bold = True
    run_gob.font.size = Pt(10)
    try: run_gob.font.color.rgb = docx.shared.RGBColor(106, 198, 79)
    except: pass
    
    run_sub = p_header.add_run("Ministerio de Educación y Derechos Humanos\nUnidad Provincial de Enlace con Universidades (UPEU)\n")
    run_sub.font.size = Pt(9.5)
    try: run_sub.font.color.rgb = docx.shared.RGBColor(100, 100, 100)
    except: pass
    
    try:
        p_line = doc.add_paragraph()
        p_line_border = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '8')
        bottom_border.set(qn('w:space'), '1')
        bottom_border.set(qn('w:color'), '007BE0')
        p_line_border.append(bottom_border)
        p_line._p.get_or_add_pPr().append(p_line_border)
    except: pass
    
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(titulo_personalizado.upper())
    run_title.font.bold = True
    run_title.font.size = Pt(15)
    try: run_title.font.color.rgb = docx.shared.RGBColor(0, 123, 224)
    except: pass
    
    p_date = doc.add_paragraph()
    aclaracion_texto = aclaracion_rango if aclaracion_rango else "Coordinación de Planificación"
    run_date = p_date.add_run(f"Fecha de emisión: {datetime.now().strftime('%d/%m/%Y')} | {aclaracion_texto}")
    run_date.font.italic = True
    run_date.font.size = Pt(9.5)
    p_date.paragraph_format.space_after = Pt(24)
    
    total_acciones = len(df)
    total_personas = df['Cantidad de personas estimadas'].sum() if total_acciones > 0 else 0
    
    table_resumen = doc.add_table(rows=2, cols=2)
    hdr_res = table_resumen.rows[0].cells
    hdr_res[0].text = "Acciones en el Reporte"
    hdr_res[1].text = "Proyección de Asistentes Global"
    for cell in hdr_res:
        set_cell_background(cell, "F5F5F5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        
    row_res = table_resumen.rows[1].cells
    row_res[0].text = str(total_acciones)
    row_res[1].text = f"{total_personas:,} personas"
    for cell in row_res:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        try: p.runs[0].font.color.rgb = docx.shared.RGBColor(0, 123, 224)
        except: pass
        
    doc.add_heading("Fichas de Planificación de Actividades", level=2)
    
    if total_acciones > 0:
        df_ordenado = df.copy()
        try:
            df_ordenado['Fecha_Limpia'] = df_ordenado['Fecha'].apply(limpiar_fecha_para_calendario)
            df_ordenado = df_ordenado.sort_values(by='Fecha_Limpia').reset_index(drop=True)
        except: pass
        
        for idx, row in df_ordenado.iterrows():
            p_act = doc.add_paragraph()
            p_act.paragraph_format.space_before = Pt(14)
            act_titulo = str(row.get('Actividad', 'Sin Nombre'))
            run_header = p_act.add_run(f"📌 Actividad {idx+1}: {act_titulo}")
            run_header.font.bold = True
            try: run_header.font.color.rgb = docx.shared.RGBColor(0, 123, 224)
            except: pass
            
            ficha_table = doc.add_table(rows=6, cols=2)
            ficha_table.style = 'Light Shading Accent 1'
            ficha_table.rows[0].cells[0].text = "Actividad:"
            ficha_table.rows[0].cells[1].text = act_titulo
            ficha_table.rows[1].cells[0].text = "Ciudad:"
            ficha_table.rows[1].cells[1].text = str(row.get('Ciudad', 'Sin especificar'))
            
            fecha_val = str(row.get('Fecha', 'Sin especificar')).strip().split(" ")[0]
            fecha_mostrar = "Sin especificar (A coordinar por el territorio)" if "sin" in fecha_val.lower() or fecha_val == "" else fecha_val
            hora_val = str(row.get('Hora', 'Sin especificar')).strip()
            hora_text = " - Sin especificar" if "sin" in hora_val.lower() or hora_val == "" else f" - {hora_val} hs"
                
            ficha_table.rows[2].cells[0].text = "Fecha:"
            ficha_table.rows[2].cells[1].text = f"{fecha_mostrar}{hora_text}"
            ficha_table.rows[3].cells[0].text = "Lugar:"
            ficha_table.rows[3].cells[1].text = str(row.get('Lugar', 'Sin especificar'))
            ficha_table.rows[4].cells[0].text = "Explicación breve de la actividad:"
            ficha_table.rows[4].cells[1].text = str(row.get('Explicación breve de la actividad', 'Sin notas adicionales'))
            ficha_table.rows[5].cells[0].text = "Cantidad de personas estimadas:"
            ficha_table.rows[5].cells[1].text = f"{int(row.get('Cantidad de personas estimadas', 0)):,} asistentes"
            
            for r_idx, r in enumerate(ficha_table.rows):
                try:
                    r.cells[0].paragraphs[0].runs[0].font.bold = True
                    if r_idx == 5:
                        r.cells[1].paragraphs[0].runs[0].font.bold = True
                        r.cells[1].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(106, 198, 79)
                except: pass
    else:
        doc.add_paragraph("No se registran actividades para los criterios de búsqueda actuales.")
        
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hashtag = p_footer.add_run("#gobiernodelosrionegrinos")
    run_hashtag.font.bold = True
    try: run_hashtag.font.color.rgb = docx.shared.RGBColor(106, 198, 79)
    except: pass
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generar_mensaje_whatsapp(df, titulo_cabecera="Agenda completa de actividades"):
    df_procesar = df.copy()
    
    # Clasificación y ordenamiento de registros
    flexibles = []
    cronologicos = []
    
    for idx, row in df_procesar.iterrows():
        fecha_val = str(row.get('Fecha', 'Sin especificar')).strip()
        if "sin especificar" in fecha_val.lower() or "a coordinar" in fecha_val.lower() or fecha_val == "":
            flexibles.append(row)
        else:
            # Adjuntamos la fecha técnica de ordenamiento de forma segura
            row_copy = row.copy()
            row_copy['_fecha_orden'] = limpiar_fecha_para_calendario(fecha_val) or "9999-12-31"
            cronologicos.append(row_copy)
            
    # Ordenar cronológicos por su fecha limpia
    if cronologicos:
        df_cron = pd.DataFrame(cronologicos).sort_values(by='_fecha_orden').reset_index(drop=True)
        cronologicos = [row for _, row in df_cron.iterrows()]
        
    # Consolidar lista de salida (flexibles primero)
    lista_final = flexibles + cronologicos
    total_eventos = len(lista_final)
    
    lines = ["🏛️ *UPEU - PLANIFICACIÓN TERRITORIAL*", f"📅 *{titulo_cabecera}*:", "─────"]
    
    if total_eventos == 0:
        lines.append("*(Sin actividades planificadas para el rango seleccionado)*")
    else:
        for idx, row in enumerate(lista_final):
            fecha_val = str(row.get('Fecha', 'Sin especificar')).strip().split(" ")[0]
            fecha_mostrar = "Sin especificar (A coordinar por el territorio)" if "sin" in fecha_val.lower() or fecha_val == "" else fecha_val
            hora_val = str(row.get('Hora', 'Sin especificar')).strip()
            hora_txt = " - Horario sin especificar" if "sin" in hora_val.lower() or hora_val == "" else f" - {hora_val} hs"
            asistencia = int(row.get('Cantidad de personas estimadas', 0))
            
            lines.append(f"📌 *Actividad {idx+1}:* {str(row.get('Actividad', 'Sin Nombre')).upper().strip()}")
            lines.append(f"📍 *Ciudad:* {str(row.get('Ciudad', 'Sin especificar')).strip()}")
            lines.append(f"📅 *Fecha:* {fecha_mostrar}{hora_txt}")
            lines.append(f"🏢 *Lugar:* {str(row.get('Lugar', 'Sin especificar')).strip()}")
            lines.append(f"🏛️ *Organismo/s involucrado/s:* {str(row.get('Organismo/Actor', 'No especificado')).strip()}")
            lines.append(f"📝 *Explicación breve de la actividad:* {str(row.get('Explicación breve de la actividad', 'Sin notas adicionales')).strip()}")
            lines.append(f"👥 *Cantidad de personas estimadas:* {f'{asistencia:,} personas estimadas' if asistencia > 0 else 'Sin especificar'}")
            
            if idx < total_eventos - 1:
                lines.append("\n❇️🔹❇️🔹❇️\n")
    lines.append("─────")
    return "\n".join(lines)

# ==========================================
# 4. DISEÑO DE LA INTERFAZ DE USUARIO (LOGO RN)
# ==========================================
if os.path.exists(LOGO_FILE): st.image(LOGO_FILE, width=180)
else: st.info("Logotipo Río Negro (SVG)")

st.markdown("---")
col_title_left, col_title_right = st.columns([4, 1.5])

with col_title_left:
    st.title("Agenda de Planificación Territorial")
    st.markdown("**Unidad Provincial de Enlace con Universidades (UPEU)** | Gobierno de Río Negro")
    st.markdown("<span class='hashtag-gestion'>#gobiernodelosrionegrinos</span>", unsafe_allow_html=True) 

with col_title_right:
    st.write("")
    st.write("")
    if st.button("🔄 Sincronizar desde GitHub", use_container_width=True):
        st.session_state.agenda = load_data()
        st.success("¡Base de datos sincronizada desde GitHub!")
        st.rerun()

if es_editor:
    tab1, tab2, tab3, tab4 = st.tabs(["🗓️ Vista de Calendario", "✍️ Carga Rápida de Actividad", "✏️ Modificar / Eliminar Actividad", "📊 Base de Datos Completa"])
else:
    tab1, tab4 = st.tabs(["🗓️ Vista de Calendario", "📊 Base de Datos Completa"])
    tab2, tab3 = None, None

# TAB 1: CALENDARIO
with tab1:
    st.header("Planificación Territorial")
    events = []
    colores_prioridad = {"ALTA": "#007BE0", "INTERMEDIA": "#333333", "BAJA": "#6AC64F"}
    
    for idx, row in st.session_state.agenda.iterrows():
        fecha_limpia = limpiar_fecha_para_calendario(row['Fecha'])
        if fecha_limpia:
            inv_val = str(row.get('Invitación a participar', '')).strip()
            tiene_inv = inv_val != "" and inv_val.lower() != "nan"
            hora_val = str(row.get('Hora', '')).strip()
            prefijo_hora = f"{hora_val} hs - " if tiene_inv and "sin" not in hora_val.lower() else ""
            act_txt = f"📍 [FECHA FLEXIBLE] {row['Actividad']}" if "sin" in str(row['Fecha']).lower() else str(row['Actividad'])
            
            titulo_mostrar = f"✉️ {prefijo_hora}[{row.get('Ciudad', 'Sin especificar')}] {act_txt}" if tiene_inv else f"{prefijo_hora}[{row.get('Ciudad', 'Sin especificar')}] {act_txt}"
            events.append({
                "title": titulo_mostrar, "start": fecha_limpia, "end": fecha_limpia, 
                "color": "#FF7A00" if tiene_inv else colores_prioridad.get(str(row['Prioridad']).upper().strip(), "#333333"),
                "extendedProps": {
                    "fecha_original": str(row['Fecha']), "hora": hora_val, "ciudad": str(row['Ciudad']), "lugar": str(row['Lugar']),
                    "explicacion": str(row['Explicación breve de la actividad']), "asistencia": int(row['Cantidad de personas estimadas']),
                    "organismo": str(row['Organismo/Actor']), "estado": str(row['Estado']), "publico": str(row['Público Destinatario']),
                    "prioridad": str(row['Prioridad']), "invitacion": inv_val if tiene_inv else "Sin invitaciones especiales"
                }
            })
            
    if len(events) > 0:
        state = calendar(events=events, options={"headerToolbar": {"left": "prev,next today", "center": "title", "right": "dayGridMonth,listMonth"}, "initialView": "dayGridMonth", "locale": "es", "buttonText": {"today": "Hoy", "month": "Mes", "list": "Lista"}}, key="calendar_agenda")
        if state.get("eventClick"):
            props = state["eventClick"]["event"].get("extendedProps", {})
            st.markdown("---")
            st.subheader("🔍 Detalle de la Actividad Seleccionada")
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                st.markdown(f"**📌 Actividad:** {state['eventClick']['event']['title']}")
                st.markdown(f"**📅 Fecha original:** `{props.get('fecha_original')}`")
                st.markdown(f"**⏰ Hora:** `{props.get('hora')}`")
                st.markdown(f"**📍 Ciudad:** {props.get('ciudad')}")
                st.markdown(f"**🏢 Lugar físico:** {props.get('lugar')}")
                st.markdown(f"**👥 Cantidad de personas estimadas:** `{props.get('asistencia')}`")
            with col_d2:
                st.markdown(f"**🚨 Prioridad:** `{props.get('prioridad')}`")
                st.markdown(f"**⚙️ Estado:** `{props.get('estado')}`")
                st.markdown(f"**✉️ Invitación a participar:** `{props.get('invitacion')}`")
                st.markdown(f"**📝 Explicación breve:** {props.get('explicacion')}")
    else: st.warning("No hay eventos programados para mostrar.")

# TAB 2: FORMULARIO DE CARGA
if es_editor and tab2 is not None:
    with tab2:
        st.header("Registrar Nueva Actividad")
        with st.form("nuevo_evento_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**📅 Fecha del Evento**")
                fecha_sin_especificar = st.checkbox("Dejar fecha sin especificar (A coordinar)", value=False)
                mes_propuesto = st.selectbox("Mes de referencia:", ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"], index=datetime.today().month-1)
                anio_propuesto = st.selectbox("Año:", [2026, 2027])
                f_fecha = st.date_input("Elegir fecha fija", datetime.today())
                st.markdown("**⏰ Horario del Evento**")
                hora_sin_especificar = st.checkbox("Dejar hora sin especificar", value=False)
                f_hora = st.time_input("Hora fija", value=time(9, 0)) 
                f_actividad = st.text_input("Nombre de la Actividad")
                f_ciudad = st.text_input("Ciudad")
                f_lugar = st.text_input("Lugar / Espacio Físico")
                f_asistencia = st.number_input("Cantidad de personas estimadas", min_value=0, step=10, value=50)
            with col2:
                f_organismo = st.text_input("Organismo / Actor principal")
                f_prioridad = st.selectbox("Nivel de Prioridad", ["ALTA", "INTERMEDIA", "BAJA"])
                f_estado = st.selectbox("Estado Actual", ["Pendiente", "En curso", "Finalizado", "Suspendido"])
                f_publico = st.text_input("Público Objetivo")
                f_invitacion = st.text_input("Invitación a participar")
                f_explicacion = st.text_area("Explicación breve de la actividad")
                
            submitted = st.form_submit_button("💾 Guardar y comitear en GitHub")
            if submitted:
                if not f_actividad or not f_ciudad: st.error("Por favor, completa 'Actividad' y 'Ciudad'.")
                else:
                    if fecha_sin_especificar:
                        meses_dict = {"Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6, "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12}
                        fecha_guardar = f"Sin especificar ({mes_propuesto} {anio_propuesto})"
                        fecha_tecnica = date(anio_propuesto, meses_dict[mes_propuesto], 1)
                    else:
                        fecha_guardar = str(f_fecha)
                        fecha_tecnica = f_fecha
                    
                    nueva = {
                        "Fecha": fecha_guardar, "Hora": "Sin especificar" if hora_sin_especificar else f_hora.strftime("%H:%M"),
                        "Semana": int(fecha_tecnica.isocalendar()[1]), "Actividad": f_actividad, "Ciudad": f_ciudad.strip(),
                        "Lugar": f_lugar.strip(), "Explicación breve de la actividad": f_explicacion, "Cantidad de personas estimadas": int(f_asistencia),
                        "Organismo/Actor": f_organismo, "Estado": f_estado, "Público Destinatario": f_publico, "Prioridad": f_prioridad, "Invitación a participar": f_invitacion
                    }
                    df_nuevo = pd.concat([st.session_state.agenda, pd.DataFrame([nueva])], ignore_index=True)
                    if push_data_to_github(df_nuevo, commit_message=f"Añadir actividad: {f_actividad}"):
                        st.session_state.agenda = df_nuevo
                        st.success("¡Actividad guardada e impactada en tu repositorio!")
                        st.rerun()

# TAB 3: MODIFICAR / ELIMINAR
if es_editor and tab3 is not None:
    with tab3:
        st.header("Editar / Cancelar Actividades")
        if len(st.session_state.agenda) > 0:
            opciones = [f"{idx} | [{row.get('Ciudad', 'Sin especificar')}] {row['Actividad']} ({row['Fecha']})" for idx, row in st.session_state.agenda.iterrows()]
            seleccionada = st.selectbox("Seleccionar Actividad a Gestionar", opciones)
            if seleccionada:
                idx_sel = int(seleccionada.split(" | ")[0])
                reg = st.session_state.agenda.loc[idx_sel]
                
                with st.form("form_edicion"):
                    col1_ed, col2_ed = st.columns(2)
                    with col1_ed:
                        ed_fecha_sin = st.checkbox("Dejar fecha sin especificar", value=("sin" in str(reg['Fecha']).lower()))
                        try: prev_f = datetime.strptime(limpiar_fecha_para_calendario(reg['Fecha']), "%Y-%m-%d").date()
                        except: prev_f = datetime.today().date()
                        ed_fecha = st.date_input("Fecha", value=prev_f)
                        ed_hora_sin = st.checkbox("Dejar hora sin especificar", value=("sin" in str(reg.get('Hora', '')).lower()))
                        try: prev_h = datetime.strptime(str(reg.get('Hora', '09:00')).strip(), "%H:%M").time()
                        except: prev_h = time(9, 0)
                        ed_hora = st.time_input("Hora", value=prev_h)
                        ed_actividad = st.text_input("Nombre de la Actividad", value=str(reg['Actividad']))
                        ed_ciudad = st.text_input("Ciudad", value=str(reg.get('Ciudad', '')))
                        ed_lugar = st.text_input("Lugar / Espacio Físico", value=str(reg.get('Lugar', '')))
                        ed_asistencia = st.number_input("Cantidad de personas estimadas", min_value=0, value=int(reg.get('Cantidad de personas estimadas', 50)))
                    with col2_ed:
                        ed_organismo = st.text_input("Organismo / Actor", value=str(reg['Organismo/Actor']))
                        ed_prioridad = st.selectbox("Prioridad", ["ALTA", "INTERMEDIA", "BAJA"], index=["ALTA", "INTERMEDIA", "BAJA"].index(str(reg['Prioridad']).upper().strip()) if str(reg['Prioridad']).upper().strip() in ["ALTA", "INTERMEDIA", "BAJA"] else 1)
                        ed_estado = st.selectbox("Estado", ["Pendiente", "En curso", "Finalizado", "Suspendido"], index=["Pendiente", "En curso", "Finalizado", "Suspendido"].index(str(reg['Estado']).capitalize().strip()) if str(reg['Estado']).capitalize().strip() in ["Pendiente", "En curso", "Finalizado", "Suspendido"] else 0)
                        ed_publico = st.text_input("Público Destinatario", value=str(reg['Público Destinatario']))
                        ed_invitacion = st.text_input("Invitación a participar", value=str(reg.get('Invitación a participar', '')))
                        ed_explicacion = st.text_area("Explicación breve", value=str(reg.get('Explicación breve de la actividad', '')))
                    
                    b_act = st.form_submit_button("🔄 Actualizar en GitHub")
                    b_eli = st.form_submit_button("❌ Eliminar del Repositorio")
                    
                    if b_act:
                        df_copia = st.session_state.agenda.copy()
                        f_g = "Sin especificar (A coordinar)" if ed_fecha_sin else str(ed_fecha)
                        f_t = date(date.today().year, date.today().month, 1) if ed_fecha_sin else ed_fecha
                        df_copia.at[idx_sel, 'Fecha'] = f_g
                        df_copia.at[idx_sel, 'Hora'] = "Sin especificar" if ed_hora_sin else ed_hora.strftime("%H:%M")
                        df_copia.at[idx_sel, 'Semana'] = int(f_t.isocalendar()[1])
                        df_copia.at[idx_sel, 'Actividad'] = ed_actividad
                        df_copia.at[idx_sel, 'Ciudad'] = ed_ciudad.strip()
                        df_copia.at[idx_sel, 'Lugar'] = ed_lugar.strip()
                        df_copia.at[idx_sel, 'Explicación breve de la actividad'] = ed_explicacion
                        df_copia.at[idx_sel, 'Cantidad de personas estimadas'] = int(ed_asistencia)
                        df_copia.at[idx_sel, 'Organismo/Actor'] = ed_organismo
                        df_copia.at[idx_sel, 'Estado'] = ed_estado
                        df_copia.at[idx_sel, 'Público Destinatario'] = ed_publico
                        df_copia.at[idx_sel, 'Prioridad'] = ed_prioridad
                        df_copia.at[idx_sel, 'Invitación a participar'] = ed_invitacion
                        if push_data_to_github(df_copia, commit_message=f"Modificar actividad: {ed_actividad}"):
                            st.session_state.agenda = df_copia
                            st.success("¡Cambios comitados con éxito!")
                            st.rerun()
                    if b_eli:
                        df_copia = st.session_state.agenda.drop(idx_sel).reset_index(drop=True)
                        if push_data_to_github(df_copia, commit_message=f"Eliminar actividad fila {idx_sel}"):
                            st.session_state.agenda = df_copia
                            st.warning("Registro removido de GitHub.")
                            st.rerun()

# TAB 4: REPORTES Y BUSCADOR
with tab4:
    st.header("Buscador y Reportes")
    df_filtrado = st.session_state.agenda.copy()
    if len(df_filtrado) > 0:
        col_f1, col_f2 = st.columns(2)
        with col_f1: search_query = st.text_input("Buscar por palabra clave...", key="search_tab4").lower()
        with col_f2:
            list_ciudades = ["Todas"] + [c for c in sorted(list(df_filtrado['Ciudad'].astype(str).str.strip().unique())) if c != 'nan' and c != '']
            filto_ciudad = st.selectbox("Filtrar por Ciudad", list_ciudades, key="filter_loc_tab4")
            
        if filto_ciudad != "Todas": df_filtrado = df_filtrado[df_filtrado['Ciudad'].str.strip() == filto_ciudad]
        if search_query:
            df_filtrado = df_filtrado[df_filtrado['Actividad'].astype(str).str.lower().str.contains(search_query) | df_filtrado['Explicación breve de la actividad'].astype(str).str.lower().str.contains(search_query)]
            
        st.dataframe(df_filtrado, use_container_width=True)
        
        st.markdown("### 📤 Generar y Exportar Documentos")
        try: semanas_disponibles = sorted([int(s) for s in df_filtrado['Semana'].dropna().unique() if str(s).strip() != "" and str(s).lower() != "nan"])
        except: semanas_disponibles = []
            
        df_filtrado['_temp_mes'] = df_filtrado['Fecha'].apply(obtener_mes_nombre)
        meses_disponibles = [m for m in ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"] if m in df_filtrado['_temp_mes'].unique()]
        df_filtrado = df_filtrado.drop(columns=['_temp_mes'])
            
        col_config, col_selectores, col_down1, col_down2 = st.columns([2.2, 1.5, 1.3, 1.3])
        with col_config:
            rango_reporte = st.radio("Alcance temporal de la descarga:", ["Agenda Completa (Historial + Futuro)", "Desde Hoy hacia adelante", "Filtrar por una semana específica", "Filtrar por un mes específico"])
            solo_con_invitacion = st.checkbox("🔍 Filtrar SOLO actividades con 'Invitación a participar'", value=False)
            
        df_descarga = df_filtrado.copy()
        titulo_word, titulo_whatsapp, rango_aclaracion_word = "REPORTE PLANIFICACION TERRITORIAL", "Cronograma de actividades", "Coordinación Interinstitucional"
        
        if rango_reporte == "Desde Hoy hacia adelante":
            hoy_d = date.today()
            df_descarga['_comp'] = df_descarga['Fecha'].apply(lambda x: datetime.strptime(limpiar_fecha_para_calendario(x), "%Y-%m-%d").date() if limpiar_fecha_para_calendario(x) else date(2000,1,1))
            df_descarga = df_descarga[df_descarga['_comp'] >= hoy_d].drop(columns=['_comp'])
            rango_aclaracion_word = "Planificación desde hoy"
        elif rango_reporte == "Agenda Completa (Historial + Futuro)":
            titulo_word, titulo_whatsapp, rango_aclaracion_word = "REPORTE COMPLETO DE ACTIVIDADES", "Agenda completa de actividades", "Historial completo de gestión"
        elif rango_reporte == "Filtrar por una semana específica" and semanas_disponibles:
            with col_selectores:
                sem_e = st.selectbox("Seleccionar Semana:", semanas_disponibles)
                df_descarga = df_descarga[df_descarga['Semana'] == sem_e]
                titulo_word, titulo_whatsapp = f"REPORTE SEMANA {sem_e}", f"Planificación - Semana {sem_e}"
        elif rango_reporte == "Filtrar por un mes específico" and meses_disponibles:
            with col_selectores:
                mes_e = st.selectbox("Seleccionar Mes:", meses_disponibles)
                df_descarga['_m_c'] = df_descarga['Fecha'].apply(obtener_mes_nombre)
                df_descarga = df_descarga[df_descarga['_m_c'] == mes_e].drop(columns=['_m_c'])
                titulo_word, titulo_whatsapp = f"REPORTE - {mes_e.upper()}", f"Planificación - Mes de {mes_e}"
                    
        if solo_con_invitacion:
            df_descarga = df_descarga[df_descarga['Invitación a participar'].notna() & (df_descarga['Invitación a participar'].astype(str).str.strip() != "")]
            titulo_word += " - PROTOCOLO"
            titulo_whatsapp += " (Protocolo)"
            
        with col_down1:
            out_e = io.BytesIO()
            with pd.ExcelWriter(out_e, engine='openpyxl') as writer: df_descarga.to_excel(writer, index=False, sheet_name='Agenda')
            st.download_button(label="📥 Descargar Excel", data=out_e.getvalue(), file_name="agenda_filtrada_territorial.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
        with col_down2:
            if LIBRERIA_DOCX_LISTA:
                st.download_button(label="📝 Descargar Reporte Word", data=crear_reporte_word_areas(df_descarga, titulo_personalizado=titulo_word, aclaracion_rango=rango_aclaracion_word), file_name="Reporte_Planificacion_Territorial.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                
        st.markdown("---")
        st.markdown("### 💬 Copiar Reporte para WhatsApp")
        st.code(generar_mensaje_whatsapp(df_descarga, titulo_cabecera=titulo_whatsapp), language="text")
    else: st.warning("La base de datos en GitHub está vacía o cargando.")
